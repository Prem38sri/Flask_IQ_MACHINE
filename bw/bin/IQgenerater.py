import json
import docx
import os
import sys

argList = sys.argv


with open("/apps/tibco/IQMachine/upload/"+argList[1],'r') as readfile:
        readfile_data = json.load(readfile)

hostname = readfile_data['hostname']
free_mem = readfile_data['free_mem']
cpu_check = readfile_data['cpu_check']
max_user_process = readfile_data['max_user_process']
stack_size = readfile_data['stack_size']
file_limit = readfile_data['file_limit']
df_apps_tibco = readfile_data['df_apps_tibco']
df_logs_tiblog = readfile_data['df_logs_tiblog']
max_user_instance = readfile_data['max_user_instance']
installed_software = readfile_data['installed_software']
bwagent_heap_size = readfile_data['bwagent_heap_size']
bwappnode_heap_size = readfile_data['bwappnode_heap_size']
logback_back = readfile_data['logback_back']
bwagent_process = readfile_data['bwagent_process']
agents_network = readfile_data['agents_network']
folder_tibco = readfile_data['folder_tibco']
folder_cft = readfile_data['folder_cft']

doc = docx.Document('/apps/tibco/IQMachine/bw/conf/IQ_template.docx')

doc.paragraphs[2].add_run(hostname)
doc.tables[2].rows[1].cells[3].text=free_mem
doc.tables[2].rows[2].cells[3].text=cpu_check
doc.tables[2].rows[3].cells[3].text=max_user_process
doc.tables[2].rows[4].cells[3].text=stack_size
doc.tables[2].rows[5].cells[3].text=file_limit
doc.tables[2].rows[6].cells[3].text=df_apps_tibco
doc.tables[2].rows[7].cells[3].text=df_logs_tiblog
doc.tables[2].rows[8].cells[3].text=max_user_instance

doc.tables[3].rows[1].cells[3].text=installed_software
doc.tables[3].rows[2].cells[3].text=bwagent_heap_size
doc.tables[3].rows[3].cells[3].text=bwappnode_heap_size
doc.tables[3].rows[4].cells[3].text=logback_back
doc.tables[3].rows[5].cells[3].text=bwagent_process
doc.tables[3].rows[6].cells[3].text=agents_network
doc.tables[3].rows[7].cells[3].text=folder_tibco
doc.tables[3].rows[8].cells[3].text=folder_cft

doc.save("/apps/tibco/IQMachine/bw/out/IQ_"+hostname+".docx")
os.system("echo \"BW IQ DOC\"\ | mailx -s \"IQ DOC BW "+hostname+"\" -a /apps/tibco/IQMachine/bw/out/IQ_"+hostname+".docx premendra.srivastava@sanofi.com")
#x = "echo \"BW IQ DOC\"|mailx -s \"IQ DOC BW "+hostname+"\" -a /apps/tibco/IQMachine/bw/out/IQ_"+hostname+".docx premendra.srivastava@sanofi.com"
#print x
#os.system("mv /apps/tibco/IQMachine/bw/out/IQ_"+hostname+".docx /apps/tibco/IQMachine/bw/tmp/IQ_"+hostname+".docx")
